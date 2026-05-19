"""Generate Access.docx – short access requirements per resource."""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- styles --
style = doc.styles["Normal"]
style.font.name = "Segoe UI"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# -- title --
t = doc.add_heading("SQL → Databricks Migration Studio", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Access Requirements — Resource Summary")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph("")  # spacer


def add_section(title, rows):
    doc.add_heading(title, level=2)
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hdrs = rows[0]
    for i, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # data
    for row in rows[1:]:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph("")  # spacer


# ── 1. Source ──
add_section("1. Source — Azure SQL", [
    ["Resource", "Access Required", "Purpose"],
    ["SQL Server", "SQL Login (username + password)", "Authenticate to server"],
    ["SQL Database", "db_datareader role", "Read tables, views, schemas"],
    ["SQL Database", "VIEW CHANGE TRACKING (optional)", "CDC / incremental loads"],
    ["SQL Firewall", "Allow Azure Services = ON", "App Service connectivity"],
])

# ── 2. Destination ──
add_section("2. Destination — Azure Databricks", [
    ["Resource", "Access Required", "Purpose"],
    ["Workspace", "Personal Access Token (PAT)", "REST API authentication"],
    ["Workspace", "Workspace Admin or Can Manage", "Upload notebooks, create jobs"],
    ["Unity Catalog", "Catalog Owner on target catalogs", "Create schemas, tables, volumes"],
    ["Compute", "Can Attach To (cluster)", "Run pipeline jobs"],
    ["SQL Warehouse", "Can Use", "DDL, metadata queries"],
])

# ── 3. Storage ──
add_section("3. Storage — ADLS Gen2", [
    ["Resource", "Access Required", "Assigned To", "Purpose"],
    ["Storage Account", "Storage Blob Data Owner", "Access Connector (MI)", "R/W landing, bronze, silver data"],
    ["Access Connector", "Storage Credential in UC", "Unity Catalog", "Bridge UC ↔ ADLS"],
    ["External Location", "Created by Catalog Owner", "Unity Catalog", "Map ADLS paths to UC"],
])

# ── 4. Deployment ──
add_section("4. Deployment — App Service + ACR", [
    ["Resource", "Access Required", "Who", "Purpose"],
    ["Subscription", "Register Microsoft.Web", "Owner", "Enable App Service provider"],
    ["Subscription", "Register Microsoft.ContainerRegistry", "Owner", "Enable ACR provider"],
    ["Resource Group", "Contributor", "Deployer", "Create ACR, Plan, Web App"],
    ["Container Registry", "AcrPush (in Contributor)", "Deployer", "Build & push Docker images"],
    ["App Service Plan", "Created by Contributor", "Deployer", "Linux B1 host"],
    ["Web App", "Created by Contributor", "Deployer", "Run Migration Studio container"],
])

# ── 5. Network ──
add_section("5. Network / Firewall", [
    ["Direction", "Port", "Protocol", "Purpose"],
    ["App Service → Azure SQL", "1433", "TCP", "Data extraction"],
    ["App Service → Databricks", "443", "HTTPS", "REST API calls"],
    ["App Service → ADLS", "443", "HTTPS", "File operations"],
    ["Browser → App Service", "443", "HTTPS", "User accesses UI"],
])

# ── Shortcut ──
doc.add_heading("Simplest Access (covers all deployment)", level=2)
p = doc.add_paragraph()
r = p.add_run("Contributor")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(37, 99, 235)
p.add_run("  on the Resource Group — covers ACR, App Plan, Web App, image push.\n")
p2 = doc.add_paragraph()
r2 = p2.add_run("⚠ Still required: ")
r2.bold = True
p2.add_run("Subscription Owner must register Microsoft.Web and Microsoft.ContainerRegistry (one-time).")

# ── Save ──
out = os.path.join(os.path.dirname(__file__), "Access.docx")
doc.save(out)
print(f"✅ Saved: {out}")
